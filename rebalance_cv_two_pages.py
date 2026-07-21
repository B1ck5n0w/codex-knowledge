from copy import deepcopy
from io import BytesIO
import os
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.shared import Pt
from lxml import etree
from PIL import Image, ImageDraw, ImageFont


SOURCE = r"C:\Users\chder\Documents\Chris Derix Privat\Christian_Derix_CV_PO_JUL_2026_layout-korrigiert.docx"
OUTPUT = r"C:\Users\chder\Documents\Chris Derix Privat\Christian_Derix_CV_PO_JUL_2026_2-Seiten-ausgewogen.docx"


def clear_runs(paragraph):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def add_run(paragraph, text, template):
    run = paragraph.add_run(text)
    if template._element.rPr is not None:
        run._element.insert(0, deepcopy(template._element.rPr))


def set_plain(paragraph, text):
    template = paragraph.runs[0]
    clear_runs(paragraph)
    add_run(paragraph, text, template)


def set_bullet(paragraph, label, description):
    bold_template = paragraph.runs[0]
    normal_template = paragraph.runs[1]
    clear_runs(paragraph)
    add_run(paragraph, f"{label}:", bold_template)
    add_run(paragraph, f" {description}", normal_template)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


doc = Document(SOURCE)
cell = doc.tables[0].cell(0, 4)

# Remove hidden generator remnants that otherwise create an empty final page.
for paragraph in list(doc.paragraphs):
    if paragraph.text == "." or paragraph.text.startswith("#HRJ#"):
        remove_paragraph(paragraph)

set_plain(
    cell.paragraphs[2],
    "Product Owner und E-Commerce-Experte mit Erfahrung in der Transformation digitaler Plattformen. "
    "Schwerpunkte sind die Übersetzung komplexer Anforderungen in umsetzbare Roadmaps und die Steuerung interdisziplinärer Teams. "
    "Expertise in Headless-Architekturen, Produktdaten, CRO und internationaler Skalierung.",
)

# The added RZH section stays complete, but is made more compact. Older roles remain intact.
replacements = {
    "Projekttransformation:": ("Projekttransformation", "Digitalisierung des Projekts „Hilfsmittel digitale Verarbeitung“."),
    "Roadmap-Steuerung:": ("Roadmap-Steuerung", "Planung, Priorisierung und Steuerung der Produkt-Roadmap."),
    "Anforderungsmanagement:": ("Anforderungsmanagement", "Überführung fachlicher Anforderungen in technische Umsetzungskonzepte."),
    "Teamkoordination:": ("Teamkoordination", "Zusammenarbeit von Projektleitung, Fachbereich und Entwicklung steuern."),
    "Backlog-Management:": ("Backlog-Management", "Anforderungen, User Stories und Akzeptanzkriterien pflegen."),
    "Stakeholder-Management & C-Level Reporting:": ("Stakeholder & Leadership", "Schnittstelle zu Fachbereichen und Geschäftsführung sowie Führung interdisziplinärer Teams und Agenturen nach SCRUM."),
    "Kampagnenmanagement:": ("Kampagnen & CRM", "Planung von Verkaufsaktionen sowie Customer-Journey- und E-Mail-Marketing-Automation."),
    "Performance Marketing:": ("Performance & Lifecycle", "SEM-Optimierung, Web-Analysen und E-Mail-Kampagnen zur Kundenbindung."),
    "UX- & Prozessoptimierung:": ("UX, Prozesse & KPIs", "Conversion- und Workflow-Optimierung sowie datenbasierte Erfolgsanalyse."),
}
for paragraph in cell.paragraphs:
    for prefix, replacement in replacements.items():
        if paragraph.text.startswith(prefix):
            set_bullet(paragraph, *replacement)
            break

# A half-point reduction keeps the dense CV readable while avoiding another page.
for paragraph in cell.paragraphs:
    if paragraph.style.name == "div_document_ul_li":
        for run in paragraph.runs:
            run.font.size = Pt(8.5)

# This point overlaps with the retained feature-development and CRO experience.
for paragraph in list(cell.paragraphs):
    if paragraph.text.startswith((
        "Agile Projektsteuerung & Leadership:",
        "Retention & CRM:",
        "Lifecycle-Marketing:",
        "Erfolgsanalyse:",
    )):
        remove_paragraph(paragraph)

doc.save(OUTPUT)

# Create a larger, cleaner handwritten-style signature.
canvas = Image.new("RGBA", (800, 112), (255, 255, 255, 0))
draw = ImageDraw.Draw(canvas)
font = ImageFont.truetype(r"C:\Windows\Fonts\FRSCRIPT.TTF", 80)
draw.text((10, -6), "Christian Derix", font=font, fill=(20, 20, 20, 255))
signature = BytesIO()
canvas.save(signature, format="PNG")

ns = {
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}
with ZipFile(OUTPUT, "r") as archive:
    relationships = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
    signature_rid = next(
        rel.get("Id") for rel in relationships
        if rel.get("Target") == "media/image9.png"
    )
    document_xml = etree.fromstring(archive.read("word/document.xml"))
    for blip in document_xml.xpath(".//a:blip", namespaces=ns):
        if blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed") == signature_rid:
            inline = blip.xpath("ancestor::wp:inline[1]", namespaces=ns)[0]
            inline.find("wp:extent", namespaces=ns).set("cx", "2142000")
            inline.find("wp:extent", namespaces=ns).set("cy", "300000")
            inline.find(".//a:xfrm/a:ext", namespaces=ns).set("cx", "2142000")
            inline.find(".//a:xfrm/a:ext", namespaces=ns).set("cy", "300000")
    document_bytes = etree.tostring(document_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

    temporary = OUTPUT + ".tmp"
    with ZipFile(temporary, "w", ZIP_DEFLATED) as updated:
        for entry in archive.infolist():
            if entry.filename == "word/document.xml":
                updated.writestr(entry, document_bytes)
            elif entry.filename == "word/media/image9.png":
                updated.writestr(entry, signature.getvalue())
            else:
                updated.writestr(entry, archive.read(entry.filename))
os.replace(temporary, OUTPUT)
